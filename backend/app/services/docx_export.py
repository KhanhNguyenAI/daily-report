"""Chuyển báo cáo markdown thành file Word (.docx)."""

import io
import re

from docx import Document
from docx.shared import Pt

TITLE = {"daily": "Daily Report", "weekly": "Weekly Report"}


def _clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text.strip()


def report_to_docx(report: dict) -> bytes:
    doc = Document()
    # Font mặc định — Word tự thay font Đông Á phù hợp cho tiếng Nhật
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    period = report.get("period_start", "")
    if report.get("period_end") and report["period_end"] != period:
        period += f" – {report['period_end']}"
    doc.add_heading(TITLE.get(report.get("type", ""), "Report"), level=0)
    doc.add_paragraph(period)

    for raw in report.get("content", "").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            text = _clean_inline(line.lstrip("#"))
            doc.add_heading(text, level=1)
        elif line.startswith(("- ", "* ", "+ ")):
            doc.add_paragraph(_clean_inline(line[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(_clean_inline(re.sub(r"^\d+\.\s*", "", line)), style="List Number")
        else:
            doc.add_paragraph(_clean_inline(line))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
